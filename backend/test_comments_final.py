#!/usr/bin/env python3
"""
Script de teste para verificar se os comentários estão sendo adicionados corretamente
no DOCX com formatação visual.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app.analysis.docx_comments import add_error_comments_to_docx
from docx import Document
import io

def create_test_docx():
    """Cria um documento DOCX de teste com conteúdo de contrato."""
    doc = Document()

    # Adiciona título
    doc.add_heading('CONTRATO DE PRESTAÇÃO DE SERVIÇOS', 0)

    # Adiciona cláusulas com erros intencionais
    doc.add_paragraph('CLÁUSULA PRIMEIRA - DO OBJETO')
    doc.add_paragraph('O presente contrato tem por objeto a prestação de serviços de consultoria em tecnologia da informação, conforme especificado no Anexo I.')

    doc.add_paragraph('CLÁUSULA SEGUNDA - DO VALOR')
    doc.add_paragraph('O valor total do contrato é de R$ 100.000,00 (cem mil reais), a ser pago em parcelas mensais de R$ 10.000,00.')

    doc.add_paragraph('CLÁUSULA TERCEIRA - DA VIGÊNCIA')
    doc.add_paragraph('O presente contrato vigorará por 12 (doze) meses, iniciando-se em __/__/____ e terminando em __/__/____.')

    doc.add_paragraph('CLÁUSULA QUARTA - DAS OBRIGAÇÕES')
    doc.add_paragraph('A CONTRATADA obriga-se a cumprir todas as obrigações constantes neste instrumento, sob pena de rescisão unilateral.')

    doc.add_paragraph('CLÁUSULA QUINTA - DO CNPJ')
    doc.add_paragraph('A CONTRATADA é a empresa XYZ Ltda, CNPJ 12.345.678/0001-XX, devidamente inscrita no CNPJ.')

    # Salva em buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def test_comments():
    """Testa a adição de comentários com erros simulados."""
    print("Criando documento de teste...")

    # Cria documento de teste
    doc_buffer = create_test_docx()

    # Erros simulados organizados por cláusula (formato esperado pela função)
    errors_by_clause = {
        "CLÁUSULA TERCEIRA - DA VIGÊNCIA": [
            {
                "id_regra": "PLACEHOLDER",
                "comentario": "Texto contém placeholders não preenchidos (__/__/____)",
                "trecho_exato": "__/__/____",
                "tipo_erro": "placeholder"
            }
        ],
        "CLÁUSULA QUARTA - DAS OBRIGAÇÕES": [
            {
                "id_regra": "CLAUSULA_UNILATERAL",
                "comentario": "Cláusula permite rescisão unilateral apenas para uma das partes",
                "trecho_exato": "rescisão unilateral",
                "tipo_erro": "clausula_unilateral"
            }
        ],
        "CLÁUSULA QUINTA - DO CNPJ": [
            {
                "id_regra": "CNPJ_INVALIDO",
                "comentario": "Formato de CNPJ inválido ou placeholder",
                "trecho_exato": "12.345.678/0001-XX",
                "tipo_erro": "cnpj_invalido"
            }
        ]
    }

    print(f"Adicionando comentários para {len(errors_by_clause)} cláusulas...")

    # Adiciona comentários
    try:
        result_buffer = add_error_comments_to_docx(doc_buffer.getvalue(), errors_by_clause)

        # Salva o resultado para inspeção
        output_path = "test_comments_output.docx"
        with open(output_path, 'wb') as f:
            f.write(result_buffer)

        print(f"Documento com comentários salvo em: {output_path}")
        print("Verifique o documento no Word para confirmar que os comentários estão visíveis")
        print("e que há sublinhados vermelhos nos trechos com erro.")

        return True

    except Exception as e:
        print(f"Erro durante o teste: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    success = test_comments()
    if success:
        print("\n✅ Teste concluído com sucesso!")
    else:
        print("\n❌ Teste falhou!")