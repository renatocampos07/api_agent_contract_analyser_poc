import io
from collections import defaultdict
from copy import deepcopy
from dataclasses import dataclass
from difflib import SequenceMatcher
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from app.analysis.doc_parser import (
    NON_BREAKING_SPACES,
    ZERO_WIDTH_CHARACTERS,
    normalize_visible_text,
    segment_document,
)

# Modulo responsavel por localizar trechos normalizados e inserir comentarios nativos em arquivos DOCX.
# Extrai o nome local de uma tag XML considerando namespaces.
def _local_name(tag: str) -> str:
    if isinstance(tag, str) and '}' in tag:
        return tag.rsplit('}', 1)[1]
    return tag


# Verifica se um nodo possui ancestrais com nomes especificos.
def _has_ancestor(node, names: set[str]) -> bool:
    parent = node
    while parent is not None:
        if _local_name(parent.tag) in names:
            return True
        parent = parent.getparent()
    return False


# Estruturas auxiliares para indexar runs e alinhar offsets normalizados com offsets originais.
@dataclass
class RunSpan:
    run: Run
    run_idx: int
    norm_start: int
    norm_end: int
    raw_start: int
    raw_end: int


@dataclass
class ParagraphIndex:
    text: str
    spans: List[RunSpan]
    norm_to_raw: List[int]
    raw_to_run: List[Tuple[int, int]]
    runs: List[Run]
    run_raw_texts: List[str]


@dataclass
class MatchLocation:
    paragraph: Paragraph
    index: ParagraphIndex
    match_start: int
    match_end: int


# Extrai o texto original de um run ignorando sugestoes inseridas e mantendo delecoes.
def _run_original_text(run) -> str:
    parts: List[str] = []
    for node in run._r.iter():
        local = _local_name(node.tag)
        if local in ('t', 'instrText'):
            if _has_ancestor(node, {'ins'}):
                continue
            parts.append(node.text or '')
        elif local == 'delText':
            parts.append(node.text or '')
    return ''.join(parts)


# Normaliza sequencias de pontos duplos mantendo o mapeamento de offsets.
def _collapse_double_dots(chars: List[str], mapping: List[int]) -> Tuple[List[str], List[int]]:
    result_chars: List[str] = []
    result_map: List[int] = []
    i = 0
    while i < len(chars):
        if (
            chars[i] == '.'
            and i + 1 < len(chars)
            and chars[i + 1] == '.'
            and (i == 0 or chars[i - 1] != '.')
            and (i + 2 >= len(chars) or chars[i + 2] != '.')
        ):
            result_chars.append('.')
            result_map.append(mapping[i])
            i += 2
            continue
        result_chars.append(chars[i])
        result_map.append(mapping[i])
        i += 1
    return result_chars, result_map


# Normaliza padroes ", ," preservando os offsets originais.
def _collapse_comma_space_comma(chars: List[str], mapping: List[int]) -> Tuple[List[str], List[int]]:
    result_chars: List[str] = []
    result_map: List[int] = []
    i = 0
    while i < len(chars):
        if (
            chars[i] == ','
            and i + 2 < len(chars)
            and chars[i + 1] == ' '
            and chars[i + 2] == ','
        ):
            result_chars.append(',')
            result_map.append(mapping[i])
            result_chars.append(' ')
            result_map.append(mapping[i + 1])
            i += 3
            continue
        result_chars.append(chars[i])
        result_map.append(mapping[i])
        i += 1
    return result_chars, result_map


# Evita a duplicacao de virgulas consecutivas sem perder o mapeamento.
def _collapse_double_commas(chars: List[str], mapping: List[int]) -> Tuple[List[str], List[int]]:
    result_chars: List[str] = []
    result_map: List[int] = []
    i = 0
    while i < len(chars):
        if chars[i] == ',' and i + 1 < len(chars) and chars[i + 1] == ',':
            result_chars.append(',')
            result_map.append(mapping[i])
            i += 2
            continue
        result_chars.append(chars[i])
        result_map.append(mapping[i])
        i += 1
    return result_chars, result_map


# Aplica a mesma normalizacao textual usada pelo parser original guardando os offsets em cada etapa.
def _normalize_with_map(text: str) -> Tuple[str, List[int]]:
    if not text:
        return '', []

    chars: List[str] = []
    mapping: List[int] = []
    i = 0
    length = len(text)
    seen_non_space = False

    while i < length:
        ch = text[i]
        if ch in ZERO_WIDTH_CHARACTERS:
            i += 1
            continue
        if ch in NON_BREAKING_SPACES:
            ch = ' '
        if ch.isspace():
            ws_start = i
            i += 1
            while i < length:
                nxt = text[i]
                if nxt in ZERO_WIDTH_CHARACTERS:
                    i += 1
                    continue
                if nxt in NON_BREAKING_SPACES:
                    nxt = ' '
                if not nxt.isspace():
                    break
                i += 1
            if seen_non_space and i < length:
                chars.append(' ')
                mapping.append(ws_start)
            continue
        chars.append(ch)
        mapping.append(i)
        seen_non_space = True
        i += 1

    if chars and chars[-1] == ' ':
        chars.pop()
        mapping.pop()

    chars, mapping = _collapse_double_dots(chars, mapping)
    chars, mapping = _collapse_comma_space_comma(chars, mapping)
    chars, mapping = _collapse_double_commas(chars, mapping)

    return ''.join(chars), mapping


# Construtor do indice de paragrafo com referencias cruzadas entre texto normalizado e runs originais.
def _build_paragraph_index(paragraph: Paragraph) -> ParagraphIndex:
    runs: List[Run] = []
    run_raw_texts: List[str] = []
    raw_to_run: List[Tuple[int, int]] = []
    raw_chars: List[str] = []

    for run in paragraph.runs:
        raw_text = _run_original_text(run)
        if not raw_text:
            continue
        run_idx = len(runs)
        runs.append(run)
        run_raw_texts.append(raw_text)
        for offset, ch in enumerate(raw_text):
            raw_chars.append(ch)
            raw_to_run.append((run_idx, offset))

    full_raw_text = ''.join(raw_chars)
    normalized_text, norm_to_raw = _normalize_with_map(full_raw_text)

    if not runs:
        return ParagraphIndex('', [], [], raw_to_run, [], [])

    norm_positions: dict[int, List[Tuple[int, int]]] = defaultdict(list)
    for norm_idx, raw_idx in enumerate(norm_to_raw):
        if raw_idx < 0 or raw_idx >= len(raw_to_run):
            continue
        run_idx, offset = raw_to_run[raw_idx]
        norm_positions[run_idx].append((norm_idx, offset))

    spans: List[RunSpan] = []
    for run_idx, run in enumerate(runs):
        positions = norm_positions.get(run_idx)
        if not positions:
            continue
        norm_start = positions[0][0]
        norm_end = positions[-1][0] + 1
        raw_start = positions[0][1]
        raw_end = positions[-1][1] + 1
        spans.append(RunSpan(run=run, run_idx=run_idx, norm_start=norm_start, norm_end=norm_end, raw_start=raw_start, raw_end=raw_end))

    return ParagraphIndex(
        text=normalized_text,
        spans=spans,
        norm_to_raw=norm_to_raw,
        raw_to_run=raw_to_run,
        runs=runs,
        run_raw_texts=run_raw_texts,
    )


# Resume o texto normalizado do paragrafo e os limites de cada run.
def _collect_run_spans(paragraph: Paragraph) -> Tuple[str, List[Tuple]]:
    index = _build_paragraph_index(paragraph)
    spans = [(span.run, span.norm_start, span.norm_end) for span in index.spans]
    return index.text, spans


# Cria um run logo apos outro reaproveitando propriedades de formato.
def _create_run_after(reference_run: Run) -> Run:
    parent = reference_run._element.getparent()
    new_r = OxmlElement('w:r')
    if reference_run._element.rPr is not None:
        new_r.append(deepcopy(reference_run._element.rPr))
    insert_at = parent.index(reference_run._element) + 1
    parent.insert(insert_at, new_r)
    new_run = Run(new_r, reference_run._parent)
    return new_run


# Fragmenta o run original em prefixo, trecho alvo e sufixo retornando o trecho selecionado.
def _slice_run(run: Run, start_offset: int, end_offset: int) -> Run:
    text = run.text or ''
    if not text:
        return run

    length = len(text)
    start = max(0, min(start_offset, length))
    end = max(start, min(end_offset, length))

    if start == 0 and end == length:
        return run

    prefix = text[:start]
    target = text[start:end]
    suffix = text[end:]

    if prefix:
        run.text = prefix
        match_run = _create_run_after(run)
        match_run.text = target
    else:
        run.text = target
        match_run = run

    if suffix:
        tail_run = _create_run_after(match_run)
        tail_run.text = suffix

    return match_run


# Transforma o intervalo normalizado encontrado em uma lista de runs prontos para receber o comentario.
def _materialize_match_runs(match: MatchLocation) -> List[Run]:
    runs: List[Run] = []
    index = match.index
    if match.match_start >= match.match_end:
        return runs

    for span in index.spans:
        if span.norm_end <= match.match_start or span.norm_start >= match.match_end:
            continue

        seg_start = max(match.match_start, span.norm_start)
        seg_end = min(match.match_end, span.norm_end)
        if seg_start >= seg_end:
            continue

        if seg_start >= len(index.norm_to_raw) or seg_end - 1 >= len(index.norm_to_raw):
            continue

        raw_start_idx = index.norm_to_raw[seg_start]
        raw_end_idx = index.norm_to_raw[seg_end - 1]

        if (
            raw_start_idx < 0
            or raw_start_idx >= len(index.raw_to_run)
            or raw_end_idx < 0
            or raw_end_idx >= len(index.raw_to_run)
        ):
            continue

        run_idx_start, offset_start = index.raw_to_run[raw_start_idx]
        run_idx_end, offset_end = index.raw_to_run[raw_end_idx]

        if run_idx_start != span.run_idx:
            offset_start = span.raw_start
        if run_idx_end != span.run_idx:
            offset_end = span.raw_end - 1

        match_run = _slice_run(span.run, offset_start, offset_end + 1)
        if match_run.text:
            runs.append(match_run)

    return runs


# Localiza o paragrafo mais similar quando a busca direta falhar.
def _find_best_paragraph(doc: Document, search_text: str) -> Optional[Paragraph]:
    target_words = set((search_text or '').lower().split())
    best_ratio = 0.0
    best_paragraph: Optional[Paragraph] = None

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        words = set(text.lower().split())
        if not words:
            continue
        ratio = len(target_words & words) / len(target_words | words)
        if ratio > best_ratio:
            best_ratio = ratio
            best_paragraph = paragraph

    return best_paragraph


# Busca o trecho normalizado dentro do paragrafo e retorna metadados da correspondencia.
def find_run_with_text(paragraph: Paragraph, search_text: str) -> Optional[MatchLocation]:
    trecho = normalize_visible_text(search_text or '')
    if not trecho:
        return None

    index = _build_paragraph_index(paragraph)
    if not index.text:
        return None

    idx = index.text.find(trecho)
    match_length = len(trecho)

    if idx == -1:
        matcher = SequenceMatcher(None, index.text, trecho, autojunk=False)
        match = matcher.find_longest_match(0, len(index.text), 0, len(trecho))
        min_size = max(3, len(trecho) // 5)
        if match.size >= min_size:
            coverage = match.size / len(trecho) if trecho else 0
            if coverage < 0.6:
                return None
            idx = match.a
            match_length = match.size
        else:
            return None

    if idx < 0:
        return None

    match_end = min(len(index.text), idx + match_length)
    if match_end <= idx:
        return None

    return MatchLocation(
        paragraph=paragraph,
        index=index,
        match_start=idx,
        match_end=match_end,
    )


# Funcao principal: aplica comentarios nos trechos fornecidos respeitando segmentacao por clausula.
def add_error_comments_to_docx(docx_content: bytes, errors_by_clause: Dict[str, List[Dict]]) -> bytes:
    doc = Document(io.BytesIO(docx_content))
    seg_map = {title: paras for title, paras in segment_document(doc)}

    for clause_title, errors in errors_by_clause.items():
        for error in errors:
            trecho_exato = normalize_visible_text(error.get('trecho_exato', ''))
            if not trecho_exato:
                continue

            candidate_paragraphs: Optional[List[Paragraph]] = seg_map.get(clause_title)
            if not candidate_paragraphs and clause_title:
                for title, paras in seg_map.items():
                    if clause_title in title:
                        candidate_paragraphs = paras
                        break
            if not candidate_paragraphs:
                candidate_paragraphs = list(doc.paragraphs)

            match_location: Optional[MatchLocation] = None
            comment_runs: List[Run] = []
            target_paragraph: Optional[Paragraph] = None

            for paragraph in candidate_paragraphs:
                location = find_run_with_text(paragraph, trecho_exato)
                if location:
                    match_location = location
                    target_paragraph = location.paragraph
                    break

            if match_location:
                comment_runs = _materialize_match_runs(match_location)

            if not comment_runs:
                target_paragraph = target_paragraph or _find_best_paragraph(doc, trecho_exato)
                if target_paragraph and target_paragraph.runs:
                    comment_runs = [target_paragraph.runs[-1]]

            if not comment_runs:
                continue

            # Registra no dicionario de erro o trecho que foi efetivamente usado
            # como ancora para o comentario no DOCX. Isso permite inspecionar no
            # playground como o texto foi localizado (trecho_marcado).
            marcado = ''
            if len(comment_runs) == 1 and comment_runs[0].text:
                marcado = comment_runs[0].text
            else:
                marcado = ''.join((run.text or '') for run in comment_runs)
            if not marcado and target_paragraph is not None:
                marcado = target_paragraph.text or ''
            error['trecho_marcado'] = marcado or None

            comment_text = f"[{error.get('id_regra', 'N/A')}] {error.get('comentario', '')}"

            try:
                doc.add_comment(comment_runs, text=comment_text, author="Revisor IA")
                for run in comment_runs:
                    run.font.underline = True
                    run.font.color.rgb = RGBColor(255, 0, 0)
            except Exception as exc:
                raise RuntimeError(
                    f"Falha ao adicionar comentario para '{trecho_exato[:50]}...'"
                ) from exc

    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    return output_buffer.getvalue()