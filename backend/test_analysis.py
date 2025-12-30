#!/usr/bin/env python3
"""
Script para testar a análise com texto de exemplo contendo placeholders.
"""

import asyncio
from app.analysis.orchestrator import run_analysis_pipeline
from app.services.storage import LocalFileStorage
import io
from docx import Document

async def test_analysis():
    # Criar um documento de teste com placeholders
    doc = Document()
    doc.add_paragraph("OBJETO: O contrato tem valor de R$xx,00 e prazo até XX/XX/XXXX.")
    doc.add_paragraph("PRAZO: A vigência é de XX meses.")
    doc.add_paragraph("PROTEÇÃO DE DADOS: Conforme LGPD.")

    # Salvar em buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    file_content = buffer.getvalue()

    # Simular storage
    storage = LocalFileStorage()

    # Executar análise
    processed_bytes, report = await run_analysis_pipeline(
        file_content=file_content,
        user_id="test_user",
        storage=storage,
        use_rag=False
    )

    print("=== RELATÓRIO DE ANÁLISE ===")
    print(f"Nome arquivo: {report.nome_arquivo}")
    print(f"Erros globais: {len(report.erros_globais)}")
    for erro in report.erros_globais:
        print(f"  - {erro.id_regra}: {erro.comentario}")

    print(f"Cláusulas analisadas: {len(report.clausulas)}")
    for clausula in report.clausulas:
        print(f"  - {clausula.titulo}: {len(clausula.erros_encontrados)} erros")
        for erro in clausula.erros_encontrados:
            print(f"    * {erro.id_regra}: {erro.comentario} (trecho: '{erro.trecho_exato}')")

if __name__ == "__main__":
    asyncio.run(test_analysis())