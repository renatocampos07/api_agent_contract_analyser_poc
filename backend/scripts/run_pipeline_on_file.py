import asyncio
import json
import io
from pathlib import Path
from docx import Document

from app.analysis.doc_parser import get_paragraph_raw_text, segment_document
from app.analysis.docx_comments import add_error_comments_to_docx
from app.services.storage import LocalFileStorage


def norm(s: str) -> str:
    return (s or '').replace('\u00A0', ' ').replace('\xa0', ' ').strip()


async def main():
    base = Path(__file__).parent.parent
    uploads = base / 'data' / 'uploads'
    processed_dir = base / 'data' / 'processed'
    processed_dir.mkdir(parents=True, exist_ok=True)

    # Escolhe o primeiro .docx em uploads
    files = list(uploads.glob('*.docx'))
    if not files:
        print('Nenhum DOCX encontrado em', uploads)
        return

    src = files[0]
    print('Processando:', src.name)

    file_bytes = src.read_bytes()

    # Em vez de invocar o LLM (que requer credenciais), vamos testar diretamente a
    # inserção de comentários nativos usando a rotina XML-aware `add_error_comments_to_docx`.
    # Montamos errors_by_clause procurando o parágrafo que contém o trecho_exato.
    doc = Document(io.BytesIO(file_bytes))

    # termo problemático informado pelo usuário
    target_trecho = 'multa compensatória devida pela CONTRATADA para a CONTRATANTE'
    target_norm = norm(target_trecho).lower()

    segmented = segment_document(doc)
    errors_by_clause = {}

    matched = False
    for title, paragraphs in segmented:
        for p in paragraphs:
            if target_norm in get_paragraph_raw_text(p).lower():
                errors_by_clause.setdefault(title, []).append({
                    'id_regra': 'R003',
                    'comentario': 'Alerta: A penalidade (multa) se aplica claramente apenas à CONTRATADA.',
                    'trecho_exato': target_trecho
                })
                matched = True
                break
        if matched:
            break

    if not matched and segmented:
        # fallback: coloca no primeiro título
        first_title = segmented[0][0]
        errors_by_clause.setdefault(first_title, []).append({
            'id_regra': 'R003',
            'comentario': 'Alerta: A penalidade (multa) se aplica claramente apenas à CONTRATADA.',
            'trecho_exato': target_trecho
        })

    # Aplica comentários via rotina especializada
    result_bytes = add_error_comments_to_docx(file_bytes, errors_by_clause)

    out_name = f"processed_{src.name}"
    out_path = processed_dir / out_name
    out_path.write_bytes(result_bytes)
    print('Arquivo processado salvo em:', out_path)

    # salva um JSON simples com errors_by_clause para referência
    report_path = processed_dir / (src.stem + '_errors_by_clause.json')
    with open(report_path, 'w', encoding='utf-8') as f:
        json.dump(errors_by_clause, f, ensure_ascii=False, indent=2)
    print('errors_by_clause salvo em:', report_path)

    # Inspeciona o docx processado em busca dos trechos exatos reportados
    doc = Document(str(out_path))

    # Para cada erro que registramos, procura onde foi inserido (run/paragraph)
    for clause_title, errors in errors_by_clause.items():
        for error in errors:
            trecho = error.get('trecho_exato', '')
            trecho_norm = norm(trecho).lower()
            print('\n---')
            print('Cláusula (chave):', clause_title)
            print("trecho_exato:", trecho)

            found = False
            for p in doc.paragraphs:
                if trecho_norm in get_paragraph_raw_text(p).lower():
                    found = True
                    print('Parágrafo encontrado:', p.text[:200])
                    # encontra run que contém início do trecho (mesma lógica de docx_comments)
                    runs_text = [norm(r.text or '') for r in p.runs]
                    combined = ''.join(runs_text)
                    import re
                    combined_search = re.sub(r"\s+", ' ', combined)
                    search_norm_spaces = re.sub(r"\s+", ' ', norm(trecho))
                    idx = combined_search.lower().find(search_norm_spaces.lower())
                    print('Índice no combined:', idx)
                    if idx >= 0:
                        acc = 0
                        for i, rt in enumerate(runs_text):
                            start = acc
                            end = acc + len(rt)
                            if end > idx:
                                run = p.runs[i]
                                print('Run contendo o início (i=', i, '):', repr(run.text))
                                try:
                                    underline = getattr(run.font, 'underline', None)
                                    color = getattr(run.font, 'color', None)
                                    rgb = None
                                    if color and hasattr(color, 'rgb'):
                                        rgb = color.rgb
                                    print('Run formatting: underline=', underline, ' color=', rgb)
                                except Exception as e:
                                    print('Erro lendo formatação do run:', e)
                                break
                            acc = end
                    else:
                        print('Índice não encontrado no texto combinado dos runs.')
                        # Imprime todos os runs do parágrafo para diagnóstico
                        print('\nRuns do parágrafo (diagnóstico):')
                        for j, rr in enumerate(p.runs):
                            try:
                                u = getattr(rr.font, 'underline', None)
                                color = getattr(rr.font, 'color', None)
                                rgb = None
                                if color and hasattr(color, 'rgb'):
                                    rgb = color.rgb
                            except Exception:
                                u = None
                                rgb = None
                            print(f"  [{j}] '{repr(rr.text)}' underline={u} color={rgb}")
                    break
            if not found:
                print('Trecho_exato não encontrado no doc processado — talvez o algoritmo inseriu fallback.')


if __name__ == '__main__':
    asyncio.run(main())
